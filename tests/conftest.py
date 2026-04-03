"""Pytest fixtures for API and synthetic document payloads."""

from __future__ import annotations

import base64
import io
import os

import fitz
import pytest
from docx import Document
from fastapi.testclient import TestClient
from PIL import Image, ImageDraw

os.environ["ENVIRONMENT"] = "test"
os.environ["API_KEY"] = "test_key"
os.environ["GROQ_API_KEY"] = "test_groq_key"

from app.main import app  # noqa: E402
from app.models.schemas import AnalysisResponse, EntitiesResponse  # noqa: E402
from app.services.pipeline import set_pipeline_instance  # noqa: E402
from app.utils.helpers import decode_base64  # noqa: E402


class MockPipeline:
    """Deterministic test pipeline used to avoid heavyweight model loading."""

    async def process(self, request):
        """Return predictable responses keyed by filename and file type."""

        try:
            content = decode_base64(request.fileBase64)
        except ValueError as exc:
            return {"status": "error", "message": str(exc)}

        if len(content) < 20:
            return {"status": "error", "message": "Could not extract readable text"}

        file_name = request.fileName.lower()
        if "technology" in file_name or request.fileType == "pdf":
            return AnalysisResponse(
                status="success",
                fileName=request.fileName,
                summary="Google, Microsoft, and NVIDIA are driving strong AI-sector growth with positive momentum.",
                entities=EntitiesResponse(
                    names=[],
                    dates=[],
                    organizations=["Google", "Microsoft", "NVIDIA"],
                    amounts=[],
                    emails=[],
                    phones=[],
                ),
                sentiment="Positive",
            )

        if "cyber" in file_name or request.fileType == "docx":
            return AnalysisResponse(
                status="success",
                fileName=request.fileName,
                summary="The incident report describes a security breach and operational impact requiring remediation.",
                entities=EntitiesResponse(names=[], dates=[], organizations=[], amounts=[], emails=[], phones=[]),
                sentiment="Negative",
            )

        return AnalysisResponse(
            status="success",
            fileName=request.fileName,
            summary="Nina Lane advanced through media and design roles across multiple organizations.",
            entities=EntitiesResponse(
                names=["Nina Lane"],
                dates=["June 2020", "March 2017", "May 2020", "2017"],
                organizations=["Brightline Agency", "Blue Horizon Media", "Parsons School of Design"],
                amounts=[],
                emails=[],
                phones=[],
            ),
            sentiment="Positive",
        )


@pytest.fixture(scope="session")
def test_client() -> TestClient:
    """Return FastAPI test client with mocked pipeline dependency."""

    set_pipeline_instance(MockPipeline())
    return TestClient(app)


@pytest.fixture
def sample_pdf_b64() -> str:
    """Create base64 for a synthetic PDF containing technology-positive text."""

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Google Microsoft NVIDIA reported strong growth and positive AI adoption.")
    payload = doc.tobytes()
    doc.close()
    return base64.b64encode(payload).decode("utf-8")


@pytest.fixture
def sample_docx_b64() -> str:
    """Create base64 for a synthetic DOCX containing cybersecurity-negative text."""

    document = Document()
    document.add_heading("Cybersecurity Incident Report", level=1)
    document.add_paragraph("A severe breach affected multiple systems and disrupted operations.")
    buffer = io.BytesIO()
    document.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("utf-8")


@pytest.fixture
def sample_image_b64() -> str:
    """Create base64 for a synthetic image containing resume-like text."""

    image = Image.new("RGB", (800, 300), color=(255, 255, 255))
    draw = ImageDraw.Draw(image)
    draw.text((20, 40), "Nina Lane | Brightline Agency | Parsons School of Design", fill=(0, 0, 0))
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG")
    return base64.b64encode(buffer.getvalue()).decode("utf-8")


@pytest.fixture
def valid_headers() -> dict[str, str]:
    """Provide valid API authentication headers."""

    return {"x-api-key": "test_key"}


@pytest.fixture
def invalid_headers() -> dict[str, str]:
    """Provide invalid API authentication headers."""

    return {"x-api-key": "wrong_key"}
